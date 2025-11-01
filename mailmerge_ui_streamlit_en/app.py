
import io
import re
import zipfile
from pathlib import Path

import pandas as pd
import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches

# ---------------------- Page config ----------------------
st.set_page_config(page_title="Mail-Merge Letters (Pro)", page_icon=None, layout="wide")

# ---------------------- Constants ----------------------
EMU_PER_INCH = 914400

# ---------------------- Helpers ----------------------
def add_ltr_paragraph(doc: Document, text: str, font_name: str, font_size_pt: int, bold: bool = False):
    """Add a LEFT-to-right paragraph (aligned LEFT) with font settings."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.bold = bold
    r.font.name = font_name
    r.font.size = Pt(font_size_pt)

def sanitize_filename(name: str) -> str:
    name = re.sub(r'[\\/*?:"<>|]+', "_", str(name)).strip()
    return name or "letter"

def replace_placeholders_dynamic(text: str, row: dict) -> str:
    """Replace any {{Field}} with row[Field] (missing => empty)."""
    def repl(m):
        key = m.group(1).strip()
        return "" if key not in row else str(row.get(key, ""))
    return re.sub(r"\{\{([^}]+)\}\}", repl, text)

def format_filename_from_pattern(pattern: str, row: dict) -> str:
    """Build filename from {Column} pattern, sanitize, ensure .docx suffix."""
    class SafeDict(dict):
        def __missing__(self, key):
            return ""
    try:
        name = pattern.format_map(SafeDict(row))
    except Exception:
        name = "letter"
    name = sanitize_filename(name)
    if not name.lower().endswith(".docx"):
        name += ".docx"
    return name

def page_width_in(section) -> float:
    """Full physical page width (inches)."""
    return float(section.page_width) / EMU_PER_INCH

def content_width_in(section) -> float:
    """Usable width (inches): page_width - margins."""
    usable_emu = section.page_width - section.left_margin - section.right_margin
    return float(usable_emu) / EMU_PER_INCH

def set_section_margins(section, left_in: float, right_in: float, top_in: float, bottom_in: float):
    section.left_margin = Inches(left_in)
    section.right_margin = Inches(right_in)
    section.top_margin = Inches(top_in)
    section.bottom_margin = Inches(bottom_in)

def add_banner_to_header_footer(doc: Document,
                                top_bytes: bytes | None, bottom_bytes: bytes | None,
                                top_path: Path | None, bottom_path: Path | None,
                                edge_to_edge: bool,
                                text_margin_in: float):
    """
    Place banners in Header (top) and Footer (bottom).
    If edge_to_edge=True → temporarily zero margins/distances, add banners at full page width,
    then restore margins for body text.
    """
    sect = doc.sections[0]

    # Save current margins/distances (Word default often ~1 inch)
    orig_left = float(sect.left_margin) / EMU_PER_INCH
    orig_right = float(sect.right_margin) / EMU_PER_INCH
    orig_top = float(sect.top_margin) / EMU_PER_INCH
    orig_bottom = float(sect.bottom_margin) / EMU_PER_INCH
    orig_header_dist = float(sect.header_distance) / EMU_PER_INCH
    orig_footer_dist = float(sect.footer_distance) / EMU_PER_INCH

    if edge_to_edge:
        # Zero everything to allow full-bleed banners
        set_section_margins(sect, 0, 0, 0, 0)
        sect.header_distance = Inches(0)
        sect.footer_distance = Inches(0)
        width_in = page_width_in(sect)
    else:
        # Respect current margins for banners
        width_in = content_width_in(sect)

    # --- Header ---
    header = sect.header
    if not header.paragraphs:
        header.add_paragraph()
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_h = hp.add_run()
    if top_bytes is not None:
        run_h.add_picture(io.BytesIO(top_bytes), width=Inches(width_in))
    elif top_path and top_path.exists():
        run_h.add_picture(str(top_path), width=Inches(width_in))

    # --- Footer ---
    footer = sect.footer
    if not footer.paragraphs:
        footer.add_paragraph()
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_f = fp.add_run()
    if bottom_bytes is not None:
        run_f.add_picture(io.BytesIO(bottom_bytes), width=Inches(width_in))
    elif bottom_path and bottom_path.exists():
        run_f.add_picture(str(bottom_path), width=Inches(width_in))

    # Restore margins for body text (either original or user-defined text margins)
    if edge_to_edge:
        # Use explicit text margins requested by user
        set_section_margins(sect, text_margin_in, text_margin_in, text_margin_in, text_margin_in)
        # Keep header/footer as close as possible to the edge
        sect.header_distance = Inches(0)
        sect.footer_distance = Inches(0)
    else:
        # Keep whatever was there originally
        set_section_margins(sect, orig_left, orig_right, orig_top, orig_bottom)
        sect.header_distance = Inches(orig_header_dist)
        sect.footer_distance = Inches(orig_footer_dist)

def build_letter_docx(
    row: dict,
    salutation: str,
    ordered_fields: list,
    body_template: str,
    font_name: str,
    font_size_pt: int,
    top_banner_bytes: bytes | None,
    bottom_banner_bytes: bytes | None,
    top_fallback_path: Path | None,
    bottom_fallback_path: Path | None,
    edge_to_edge: bool,
    text_margin_in: float,
):
    """
    Create a DOCX letter (LTR, left-aligned) and return (docx_bytes, final_body_text).
    Banners are placed in header/footer and can stretch to full page width,
    while the body text keeps the requested margins.
    """
    doc = Document()

    # Banners (header/footer) + margins handling
    add_banner_to_header_footer(
        doc,
        top_banner_bytes, bottom_banner_bytes,
        top_fallback_path, bottom_fallback_path,
        edge_to_edge=edge_to_edge,
        text_margin_in=text_margin_in,
    )

    # Salutation
    add_ltr_paragraph(doc, salutation, font_name, font_size_pt, bold=True)

    # Ordered header fields under salutation
    for col in ordered_fields:
        val = row.get(col, "")
        if pd.isna(val):
            val = ""
        add_ltr_paragraph(doc, str(val), font_name, font_size_pt)

    # Spacer
    doc.add_paragraph("")

    # Body with dynamic placeholders
    body_filled = replace_placeholders_dynamic(body_template, row)
    for line in body_filled.splitlines():
        add_ltr_paragraph(doc, line, font_name, font_size_pt)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read(), body_filled

# ---------------------- UI Header ----------------------
st.markdown(
    """
    <div style="background:linear-gradient(120deg,#EFF6FF,#FFFFFF);border:1px solid #e6ecff;padding:16px 18px;border-radius:14px;margin-bottom:10px;">
      <h1 style="margin:0;font-size:1.6rem;">Mail-Merge Letters (Pro)</h1>
      <p style="margin:6px 0 0;color:#334155;">Edge-to-edge banners, with body text margins preserved.</p>
    </div>
    """,
    unsafe_allow_html=True,
)

# ---------------------- Sidebar: Settings ----------------------
with st.sidebar:
    st.header("Settings")
    font_name = st.selectbox("Font", ["Arial", "Rubik", "David", "Narkisim"], index=0)
    font_size = st.number_input("Font size (pt)", min_value=8, max_value=24, value=12, step=1)
    salutation = st.text_input("Salutation line", value="To,")
    st.markdown("---")

    st.subheader("Banners (Header/Footer)")
    st.caption("Upload images or place files named 'upper banner.png' / 'low banner.png' next to app.py")
    top_banner_file = st.file_uploader("Top banner (PNG/JPG)", type=["png", "jpg", "jpeg"], key="top_banner")
    bottom_banner_file = st.file_uploader("Bottom banner (PNG/JPG)", type=["png", "jpg", "jpeg"], key="bottom_banner")

    edge_to_edge = st.checkbox("Edge-to-edge banners (full width, beyond margins)", value=True,
                               help="Banners use the full page width. Text keeps its own margins.")
    text_margin_in = st.slider("Body text margin (inches)", min_value=0.25, max_value=2.0, value=1.0, step=0.25)

    app_dir = Path.cwd()
    top_fallback = app_dir / "upper banner.png"
    bottom_fallback = app_dir / "low banner.png"

# ---------------------- Main Layout ----------------------
left, right = st.columns([0.55, 0.45], gap="large")

DEFAULT_TEMPLATES = {
    "blue": "Hello {{FullName}},\n\nWe are delighted to invite you to our upcoming event at {{Institution}}.\nWe'd be honored to see you there.\n\nWarm regards,\nEvent Team",
    "green": "Hello {{FullName}},\n\nYou are part of Group Green. The event will take place at: {{Address}}.\nPlease confirm your attendance.\n\nBest,\nEvent Team",
    "yellow": "Hello {{FullName}},\n\nWe look forward to hosting you. Our representatives from {{Institution}} will be available for questions.\nSee you soon!\n\nBest regards,\nEvent Team",
}

with left:
    st.subheader("1) Upload your Excel")
    excel_file = st.file_uploader("Choose the guest list (XLSX)", type=["xlsx"], help="Any columns are allowed.")
    df = None
    if excel_file:
        try:
            df = pd.read_excel(excel_file)
            st.success(f"Loaded Excel with {len(df)} rows and {len(df.columns)} columns.")
            st.dataframe(df.head(10), use_container_width=True)
        except Exception as e:
            st.error(f"Failed to read Excel: {e}")

    st.subheader("2) Group & Templates")
    if df is not None:
        cols = list(df.columns)
        group_col = st.selectbox("Select the Group column", options=cols, index=(cols.index("Group") if "Group" in cols else 0))

        uniques = list(map(str, sorted(set(df[group_col].dropna().astype(str)))))[:10]
        default_blue = "כחול" if "כחול" in uniques else (uniques[0] if uniques else "כחול")
        default_green = "ירוק" if "ירוק" in uniques else (uniques[1] if len(uniques) > 1 else "ירוק")
        default_yellow = "צהוב" if "צהוב" in uniques else (uniques[2] if len(uniques) > 2 else "צהוב")

        blue_value = st.text_input("Value for Blue group", value=default_blue)
        green_value = st.text_input("Value for Green group", value=default_green)
        yellow_value = st.text_input("Value for Yellow group", value=default_yellow)

        st.markdown("Templates (placeholders allowed: {{AnyColumn}})")
        tab_b, tab_g, tab_y = st.tabs(["Blue", "Green", "Yellow"])
        with tab_b:
            tpl_blue = st.text_area("Blue template", value=DEFAULT_TEMPLATES["blue"], height=220)
        with tab_g:
            tpl_green = st.text_area("Green template", value=DEFAULT_TEMPLATES["green"], height=220)
        with tab_y:
            tpl_yellow = st.text_area("Yellow template", value=DEFAULT_TEMPLATES["yellow"], height=220)

with right:
    st.subheader("3) Header fields & filename")
    if df is not None:
        all_cols = list(df.columns)
        header_fields = st.multiselect(
            "Select the fields (order matters) to appear under the salutation:",
            options=all_cols,
            default=[c for c in ["FullName", "Address", "Institution"] if c in all_cols],
            help="The order you select will be the order in the document header block.",
        )

        filename_pattern = st.text_input("Filename pattern ('.docx' added if missing)", value="{Full Name} - {Country}")
        st.caption("Use {ColumnName} placeholders, e.g. {FullName}, {Group}. Unknown names become empty.")

        st.subheader("4) Preview (single row)")
        options = list(range(len(df)))
        if "FullName" in df.columns:
            labels = [f"{i}: {df.iloc[i]['FullName']}" for i in options]
            pick = st.selectbox("Pick a row to preview", options=options, format_func=lambda i: labels[i])
        else:
            pick = st.number_input("Pick a row index", min_value=0, max_value=max(0, len(df) - 1), value=0, step=1)

        row = df.iloc[int(pick)].to_dict()
        gval = str(row.get(group_col, "") )

        if gval == str(blue_value):
            tpl = tpl_blue
        elif gval == str(green_value):
            tpl = tpl_green
        elif gval == str(yellow_value):
            tpl = tpl_yellow
        else:
            tpl = tpl_blue  # fallback

        file_name = format_filename_from_pattern(filename_pattern, row)
        content, body_filled = build_letter_docx(
            row=row,
            salutation=salutation,
            ordered_fields=header_fields,
            body_template=tpl,
            font_name=font_name,
            font_size_pt=font_size,
            top_banner_bytes=(top_banner_file.read() if top_banner_file else None),
            bottom_banner_bytes=(bottom_banner_file.read() if bottom_banner_file else None),
            top_fallback_path=top_fallback,
            bottom_fallback_path=bottom_fallback,
            edge_to_edge=edge_to_edge,
            text_margin_in=text_margin_in,
        )

        st.write("Filename example:", file_name)
        st.write("Header fields:", ", ".join(header_fields) if header_fields else "— none —")
        st.text_area("Body after placeholders:", value=body_filled, height=220)
        st.download_button(
            "Download preview DOCX",
            data=content,
            file_name=file_name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    st.subheader("5) Generate all")
    gen = st.button("Create all letters and download ZIP", use_container_width=True)
    if gen:
        if df is None:
            st.error("Please upload an Excel file first.")
            st.stop()
        if group_col not in df.columns:
            st.error("Group column selection is invalid.")
            st.stop()

        zip_buffer = io.BytesIO()
        created = 0
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as z:
            for _, r in df.iterrows():
                row = r.to_dict()
                gval = str(row.get(group_col, ""))

                if gval == str(blue_value):
                    tpl = tpl_blue
                elif gval == str(green_value):
                    tpl = tpl_green
                elif gval == str(yellow_value):
                    tpl = tpl_yellow
                else:
                    st.warning(f"Skipping '{row.get('FullName','')}' – unknown group value: {gval}")
                    continue

                file_name = format_filename_from_pattern(filename_pattern, row)
                content, _ = build_letter_docx(
                    row=row,
                    salutation=salutation,
                    ordered_fields=header_fields,
                    body_template=tpl,
                    font_name=font_name,
                    font_size_pt=font_size,
                    top_banner_bytes=(top_banner_file.read() if top_banner_file else None),
                    bottom_banner_bytes=(bottom_banner_file.read() if bottom_banner_file else None),
                    top_fallback_path=top_fallback,
                    bottom_fallback_path=bottom_fallback,
                    edge_to_edge=edge_to_edge,
                    text_margin_in=text_margin_in,
                )
                z.writestr(file_name, content)
                created += 1

        zip_buffer.seek(0)
        st.success(f"Created {created} letters.")
        st.download_button(
            "Download ZIP",
            data=zip_buffer.getvalue(),
            file_name="letters_output.zip",
            mime="application/zip",
            use_container_width=True,
        )
